from docx import Document
doc=Document(r"H:\exercise\test1.docx")
print(doc.paragraphs)
paragraph=doc.paragraphs[0]
runs=paragraph.runs
print(runs)
for run in paragraph.runs:
    print(run.text)
paragraph=doc.paragraphs[1]
runs=paragraph.runs
print(runs)
for run in paragraph.runs:
    print(run.text)