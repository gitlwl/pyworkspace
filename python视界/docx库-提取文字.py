from docx import Document
#提取文字
doc=Document(r"H:\exercise\test1.docx")
print(doc.paragraphs)
for paragraph in doc.paragraphs:
    print(paragraph.text)

#提取文字块儿
